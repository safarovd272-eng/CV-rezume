"""
CV Generator - Europa Pass International Style
PDF va DOCX formatlarida chiroyli CV yaratadi
"""

import os
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import mm, cm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    HRFlowable, Image, KeepTogether
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont


# ─── Europa Pass Color Palette ────────────────────────────────────────────────
BLUE_DARK   = colors.HexColor('#003399')   # EU Blue (sidebar)
BLUE_MEDIUM = colors.HexColor('#0050A0')   # Section headers
BLUE_LIGHT  = colors.HexColor('#E8F0FB')   # Background tint
GOLD        = colors.HexColor('#FFCC00')   # EU Stars gold
WHITE       = colors.white
GRAY_DARK   = colors.HexColor('#333333')
GRAY_MED    = colors.HexColor('#666666')
GRAY_LIGHT  = colors.HexColor('#F5F5F5')
BLACK       = colors.black


def parse_education(edu_list):
    result = []
    for item in edu_list:
        parts = [p.strip() for p in item.split('|')]
        if len(parts) >= 3:
            result.append({
                'degree': parts[0] if len(parts) > 0 else '',
                'institution': parts[1] if len(parts) > 1 else '',
                'years': parts[2] if len(parts) > 2 else '',
                'gpa': parts[3] if len(parts) > 3 else '',
            })
    return result


def parse_work(work_list):
    result = []
    for item in work_list:
        parts = [p.strip() for p in item.split('|')]
        if len(parts) >= 3:
            result.append({
                'position': parts[0] if len(parts) > 0 else '',
                'company': parts[1] if len(parts) > 1 else '',
                'years': parts[2] if len(parts) > 2 else '',
                'description': parts[3] if len(parts) > 3 else '',
            })
    return result


def parse_skills(skills_list):
    result = []
    for item in skills_list:
        if ':' in item:
            cat, skills = item.split(':', 1)
            result.append({'category': cat.strip(), 'skills': skills.strip()})
    return result


def parse_languages(lang_list):
    result = []
    for item in lang_list:
        parts = [p.strip() for p in item.split('|')]
        if len(parts) >= 2:
            result.append({'language': parts[0], 'level': parts[1]})
    return result


def parse_certificates(cert_list):
    result = []
    for item in cert_list:
        parts = [p.strip() for p in item.split('|')]
        if len(parts) >= 2:
            result.append({
                'name': parts[0],
                'org': parts[1] if len(parts) > 1 else '',
                'year': parts[2] if len(parts) > 2 else '',
            })
    return result


def get_level_bar(level_text):
    """Convert language level to dots (CEFR scale)"""
    level_map = {
        'a1': 1, 'a2': 2, 'b1': 3, 'b2': 4, 'c1': 5, 'c2': 6,
        'native': 6, 'ona tili': 6, 'родной': 6,
        'beginner': 1, 'elementary': 2, 'intermediate': 3,
        'upper intermediate': 4, 'advanced': 5, 'proficient': 6,
    }
    key = level_text.strip().lower()
    dots = level_map.get(key, 3)
    filled = '●' * dots + '○' * (6 - dots)
    return filled


# ─── PDF Generation ───────────────────────────────────────────────────────────

def generate_pdf(data: dict, output_path: str):
    """Europa Pass style PDF yaratadi"""
    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        leftMargin=0,
        rightMargin=0,
        topMargin=0,
        bottomMargin=0,
    )

    page_width, page_height = A4
    sidebar_width = 68 * mm
    main_width = page_width - sidebar_width

    story = []

    # ── Custom Styles ─────────────────────────────────────────────────────────
    name_style = ParagraphStyle(
        'NameStyle',
        fontName='Helvetica-Bold',
        fontSize=22,
        textColor=WHITE,
        spaceAfter=2,
        leading=26,
    )
    title_style = ParagraphStyle(
        'TitleStyle',
        fontName='Helvetica',
        fontSize=12,
        textColor=GOLD,
        spaceAfter=0,
        leading=16,
    )
    sidebar_label_style = ParagraphStyle(
        'SidebarLabel',
        fontName='Helvetica-Bold',
        fontSize=7,
        textColor=GOLD,
        spaceBefore=8,
        spaceAfter=1,
        leading=9,
    )
    sidebar_value_style = ParagraphStyle(
        'SidebarValue',
        fontName='Helvetica',
        fontSize=8,
        textColor=WHITE,
        spaceAfter=0,
        leading=10,
    )
    sidebar_section_style = ParagraphStyle(
        'SidebarSection',
        fontName='Helvetica-Bold',
        fontSize=9,
        textColor=GOLD,
        spaceBefore=14,
        spaceAfter=4,
        leading=11,
        borderPad=0,
    )
    section_title_style = ParagraphStyle(
        'SectionTitle',
        fontName='Helvetica-Bold',
        fontSize=11,
        textColor=BLUE_MEDIUM,
        spaceBefore=14,
        spaceAfter=4,
        leading=14,
    )
    normal_style = ParagraphStyle(
        'NormalCV',
        fontName='Helvetica',
        fontSize=9,
        textColor=GRAY_DARK,
        spaceAfter=2,
        leading=12,
    )
    bold_style = ParagraphStyle(
        'BoldCV',
        fontName='Helvetica-Bold',
        fontSize=9,
        textColor=GRAY_DARK,
        spaceAfter=1,
        leading=12,
    )
    italic_style = ParagraphStyle(
        'ItalicCV',
        fontName='Helvetica-Oblique',
        fontSize=8,
        textColor=GRAY_MED,
        spaceAfter=2,
        leading=10,
    )
    objective_style = ParagraphStyle(
        'ObjectiveStyle',
        fontName='Helvetica',
        fontSize=9,
        textColor=GRAY_DARK,
        spaceAfter=4,
        leading=13,
        alignment=TA_JUSTIFY,
    )

    # ── Sidebar Content ───────────────────────────────────────────────────────
    sidebar_content = []

    # Photo
    photo_path = data.get('photo')
    if photo_path and os.path.exists(photo_path):
        try:
            from PIL import Image as PILImage
            img = PILImage.open(photo_path)
            w, h = img.size
            ratio = w / h
            img_w = sidebar_width - 20 * mm
            img_h = img_w / ratio
            img_obj = Image(photo_path, width=img_w, height=min(img_h, 50 * mm))
            img_obj.hAlign = 'CENTER'
            sidebar_content.append(Spacer(1, 8 * mm))
            sidebar_content.append(img_obj)
            sidebar_content.append(Spacer(1, 4 * mm))
        except:
            sidebar_content.append(Spacer(1, 60 * mm))
    else:
        # Placeholder circle
        sidebar_content.append(Spacer(1, 8 * mm))
        initials = f"{data.get('first_name', '')[:1]}{data.get('last_name', '')[:1]}".upper()
        init_style = ParagraphStyle(
            'Initials', fontName='Helvetica-Bold', fontSize=28,
            textColor=GOLD, alignment=TA_CENTER
        )
        # Draw a circle placeholder
        circle_table = Table(
            [[Paragraph(initials, init_style)]],
            colWidths=[44 * mm], rowHeights=[44 * mm]
        )
        circle_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#0050A0')),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('ROUNDEDCORNERS', [22]),
            ('BOX', (0, 0), (-1, -1), 2, GOLD),
        ]))
        sidebar_content.append(circle_table)
        sidebar_content.append(Spacer(1, 4 * mm))

    # Contact info
    sidebar_content.append(Paragraph("CONTACT", sidebar_section_style))
    sidebar_content.append(HRFlowable(width=sidebar_width - 16 * mm, thickness=0.5, color=GOLD, spaceAfter=4))

    contact_items = [
        ("📧", data.get('email', '')),
        ("📱", data.get('phone', '')),
        ("📍", data.get('address', '')),
    ]
    for icon, value in contact_items:
        if value:
            sidebar_content.append(Paragraph(f"{icon} {value}", sidebar_value_style))

    if data.get('linkedin'):
        sidebar_content.append(Paragraph(f"in  {data['linkedin']}", sidebar_value_style))
    if data.get('github'):
        sidebar_content.append(Paragraph(f"⌨  {data['github']}", sidebar_value_style))
    if data.get('website'):
        sidebar_content.append(Paragraph(f"🌐 {data['website']}", sidebar_value_style))

    # Personal info
    sidebar_content.append(Paragraph("PERSONAL", sidebar_section_style))
    sidebar_content.append(HRFlowable(width=sidebar_width - 16 * mm, thickness=0.5, color=GOLD, spaceAfter=4))

    personal_items = [
        ("Date of Birth", data.get('dob', '')),
        ("Nationality", data.get('nationality', '')),
    ]
    for label, val in personal_items:
        if val:
            sidebar_content.append(Paragraph(label.upper(), sidebar_label_style))
            sidebar_content.append(Paragraph(val, sidebar_value_style))

    # Languages
    lang_list = parse_languages(data.get('lang_list', []))
    if lang_list:
        sidebar_content.append(Paragraph("LANGUAGES", sidebar_section_style))
        sidebar_content.append(HRFlowable(width=sidebar_width - 16 * mm, thickness=0.5, color=GOLD, spaceAfter=4))
        for lang in lang_list:
            bar = get_level_bar(lang['level'])
            lname_style = ParagraphStyle('LangName', fontName='Helvetica-Bold', fontSize=8, textColor=WHITE, leading=10)
            llevel_style = ParagraphStyle('LangLevel', fontName='Helvetica', fontSize=7, textColor=GOLD, leading=9)
            sidebar_content.append(Paragraph(lang['language'], lname_style))
            sidebar_content.append(Paragraph(f"{lang['level']}  {bar}", llevel_style))

    # Skills in sidebar
    skills_list = parse_skills(data.get('skills_list', []))
    if skills_list:
        sidebar_content.append(Paragraph("SKILLS", sidebar_section_style))
        sidebar_content.append(HRFlowable(width=sidebar_width - 16 * mm, thickness=0.5, color=GOLD, spaceAfter=4))
        for sk in skills_list:
            cat_style = ParagraphStyle('SkCat', fontName='Helvetica-Bold', fontSize=7, textColor=GOLD, leading=9, spaceBefore=4)
            sk_style = ParagraphStyle('SkVal', fontName='Helvetica', fontSize=7.5, textColor=WHITE, leading=9)
            sidebar_content.append(Paragraph(sk['category'].upper(), cat_style))
            sidebar_content.append(Paragraph(sk['skills'], sk_style))

    # ── Main Content ──────────────────────────────────────────────────────────
    main_content = []

    # Header background done via table below
    first_name = data.get('first_name', '')
    last_name = data.get('last_name', '')

    main_content.append(Spacer(1, 4 * mm))
    main_content.append(Paragraph(f"{first_name.upper()} {last_name.upper()}", name_style))

    # Objective as subtitle
    objective = data.get('objective', '')
    if objective:
        main_content.append(Spacer(1, 2 * mm))
        main_content.append(Paragraph(objective, objective_style))

    # Education
    edu_list = parse_education(data.get('education_list', []))
    if edu_list:
        main_content.append(Paragraph("EDUCATION", section_title_style))
        main_content.append(HRFlowable(
            width=main_width - 24 * mm, thickness=1, color=BLUE_MEDIUM, spaceAfter=4
        ))
        for edu in edu_list:
            row = Table(
                [[
                    Paragraph(f"<b>{edu['degree']}</b>", bold_style),
                    Paragraph(edu['years'], italic_style)
                ]],
                colWidths=[(main_width - 30 * mm) * 0.7, (main_width - 30 * mm) * 0.3]
            )
            row.setStyle(TableStyle([('ALIGN', (1, 0), (1, 0), 'RIGHT')]))
            main_content.append(row)
            main_content.append(Paragraph(edu['institution'], italic_style))
            if edu.get('gpa'):
                main_content.append(Paragraph(f"GPA: {edu['gpa']}", italic_style))
            main_content.append(Spacer(1, 3 * mm))

    # Work Experience
    work_list = parse_work(data.get('work_list', []))
    if work_list:
        main_content.append(Paragraph("WORK EXPERIENCE", section_title_style))
        main_content.append(HRFlowable(
            width=main_width - 24 * mm, thickness=1, color=BLUE_MEDIUM, spaceAfter=4
        ))
        for job in work_list:
            row = Table(
                [[
                    Paragraph(f"<b>{job['position']}</b>", bold_style),
                    Paragraph(job['years'], italic_style)
                ]],
                colWidths=[(main_width - 30 * mm) * 0.7, (main_width - 30 * mm) * 0.3]
            )
            row.setStyle(TableStyle([('ALIGN', (1, 0), (1, 0), 'RIGHT')]))
            main_content.append(row)
            main_content.append(Paragraph(job['company'], italic_style))
            if job.get('description'):
                for desc_line in job['description'].split(','):
                    main_content.append(Paragraph(f"• {desc_line.strip()}", normal_style))
            main_content.append(Spacer(1, 3 * mm))

    # Certificates
    cert_list = parse_certificates(data.get('cert_list', []))
    if cert_list:
        main_content.append(Paragraph("CERTIFICATES", section_title_style))
        main_content.append(HRFlowable(
            width=main_width - 24 * mm, thickness=1, color=BLUE_MEDIUM, spaceAfter=4
        ))
        for cert in cert_list:
            row = Table(
                [[
                    Paragraph(f"<b>{cert['name']}</b> – {cert['org']}", bold_style),
                    Paragraph(cert['year'], italic_style)
                ]],
                colWidths=[(main_width - 30 * mm) * 0.75, (main_width - 30 * mm) * 0.25]
            )
            row.setStyle(TableStyle([('ALIGN', (1, 0), (1, 0), 'RIGHT')]))
            main_content.append(row)

    # Hobbies
    hobbies = data.get('hobbies', '')
    if hobbies:
        main_content.append(Paragraph("INTERESTS", section_title_style))
        main_content.append(HRFlowable(
            width=main_width - 24 * mm, thickness=1, color=BLUE_MEDIUM, spaceAfter=4
        ))
        main_content.append(Paragraph(hobbies, normal_style))

    # ── Build Two-Column Layout ───────────────────────────────────────────────
    # Sidebar column
    sidebar_table_data = []
    for item in sidebar_content:
        sidebar_table_data.append([item])

    sidebar_col = Table(
        sidebar_table_data,
        colWidths=[sidebar_width - 8 * mm]
    )
    sidebar_col.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), BLUE_DARK),
        ('LEFTPADDING', (0, 0), (-1, -1), 8 * mm),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4 * mm),
        ('TOPPADDING', (0, 0), (-1, -1), 0),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
    ]))

    # Main column
    main_table_data = []
    for item in main_content:
        main_table_data.append([item])

    main_col = Table(
        main_table_data,
        colWidths=[main_width - 16 * mm]
    )
    main_col.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), WHITE),
        ('LEFTPADDING', (0, 0), (-1, -1), 8 * mm),
        ('RIGHTPADDING', (0, 0), (-1, -1), 8 * mm),
        ('TOPPADDING', (0, 0), (-1, -1), 0),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
    ]))

    # Combine sidebar + main
    layout = Table(
        [[sidebar_col, main_col]],
        colWidths=[sidebar_width, main_width],
        rowHeights=[page_height]
    )
    layout.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
        ('RIGHTPADDING', (0, 0), (-1, -1), 0),
        ('TOPPADDING', (0, 0), (-1, -1), 0),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
    ]))

    story.append(layout)
    doc.build(story)


# ─── DOCX Generation ──────────────────────────────────────────────────────────

def generate_docx(data: dict, output_path: str):
    """Europa Pass style DOCX yaratadi"""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import copy

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0)
        section.left_margin = Cm(0)
        section.right_margin = Cm(0)
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)

    # Color constants
    EU_BLUE = RGBColor(0, 51, 153)
    EU_BLUE_MED = RGBColor(0, 80, 160)
    EU_GOLD = RGBColor(255, 204, 0)
    WHITE_COLOR = RGBColor(255, 255, 255)
    DARK_GRAY = RGBColor(51, 51, 51)
    MED_GRAY = RGBColor(102, 102, 102)

    def set_cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    def add_run_with_style(para, text, bold=False, italic=False,
                           font_size=10, color=None, font_name='Calibri'):
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(font_size)
        run.font.name = font_name
        if color:
            run.font.color.rgb = color
        return run

    # ── 2-column table: Sidebar | Main ────────────────────────────────────────
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.allow_autofit = False

    sidebar_cell = tbl.cell(0, 0)
    main_cell = tbl.cell(0, 1)

    # Widths
    sidebar_cell.width = Cm(6.8)
    main_cell.width = Cm(14.2)

    # Background colors
    set_cell_bg(sidebar_cell, '003399')
    set_cell_bg(main_cell, 'FFFFFF')

    # Cell vertical alignment
    sidebar_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    main_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # ── Fill Sidebar ──────────────────────────────────────────────────────────
    def sidebar_para(text, bold=False, color=WHITE_COLOR, size=9, space_before=0, space_after=2):
        p = sidebar_cell.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run.font.name = 'Calibri'
        return p

    def sidebar_section_header(text):
        p = sidebar_para(f"◆ {text}", bold=True, color=EU_GOLD, size=9, space_before=10, space_after=3)
        return p

    # Initials circle (text-based)
    initials = f"{data.get('first_name', '')[:1]}{data.get('last_name', '')[:1]}".upper()
    p_init = sidebar_cell.add_paragraph()
    p_init.paragraph_format.space_before = Pt(20)
    p_init.paragraph_format.space_after = Pt(4)
    p_init.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_init = p_init.add_run(initials)
    run_init.font.size = Pt(32)
    run_init.font.bold = True
    run_init.font.color.rgb = EU_GOLD

    # Contact
    sidebar_section_header("CONTACT")
    if data.get('email'):
        sidebar_para(f"✉  {data['email']}", size=8)
    if data.get('phone'):
        sidebar_para(f"📱 {data['phone']}", size=8)
    if data.get('address'):
        sidebar_para(f"📍 {data['address']}", size=8)
    if data.get('linkedin'):
        sidebar_para(f"in  {data['linkedin']}", size=7.5)
    if data.get('github'):
        sidebar_para(f"⌨  {data['github']}", size=7.5)
    if data.get('website'):
        sidebar_para(f"🌐 {data['website']}", size=7.5)

    # Personal
    sidebar_section_header("PERSONAL")
    if data.get('dob'):
        sidebar_para("DATE OF BIRTH", bold=True, color=EU_GOLD, size=7)
        sidebar_para(data['dob'], size=8)
    if data.get('nationality'):
        sidebar_para("NATIONALITY", bold=True, color=EU_GOLD, size=7)
        sidebar_para(data['nationality'], size=8)

    # Languages
    lang_list = parse_languages(data.get('lang_list', []))
    if lang_list:
        sidebar_section_header("LANGUAGES")
        for lang in lang_list:
            bar = get_level_bar(lang['level'])
            sidebar_para(lang['language'], bold=True, size=8.5)
            sidebar_para(f"{lang['level']}  {bar}", color=EU_GOLD, size=7.5)

    # Skills
    skills_list = parse_skills(data.get('skills_list', []))
    if skills_list:
        sidebar_section_header("SKILLS")
        for sk in skills_list:
            sidebar_para(sk['category'].upper(), bold=True, color=EU_GOLD, size=7.5)
            sidebar_para(sk['skills'], size=8)

    # ── Fill Main Content ─────────────────────────────────────────────────────
    main_cell.paragraphs[0].clear()

    def main_para(text='', bold=False, italic=False, size=10,
                  color=DARK_GRAY, space_before=0, space_after=2,
                  align=WD_ALIGN_PARAGRAPH.LEFT, indent=0.5):
        p = main_cell.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.left_indent = Cm(indent)
        p.alignment = align
        if text:
            run = p.add_run(text)
            run.bold = bold
            run.italic = italic
            run.font.size = Pt(size)
            run.font.color.rgb = color
            run.font.name = 'Calibri'
        return p

    def section_header(text):
        p = main_para(text, bold=True, size=12, color=EU_BLUE_MED, space_before=12, space_after=2)
        # Add bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '0050A0')
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    # Name
    name_p = main_cell.paragraphs[0]
    name_p.paragraph_format.space_before = Pt(16)
    name_p.paragraph_format.space_after = Pt(4)
    name_p.paragraph_format.left_indent = Cm(0.5)
    name_run = name_p.add_run(
        f"{data.get('first_name', '').upper()} {data.get('last_name', '').upper()}"
    )
    name_run.font.size = Pt(24)
    name_run.font.bold = True
    name_run.font.color.rgb = EU_BLUE
    name_run.font.name = 'Calibri'

    # Objective
    if data.get('objective'):
        main_para(data['objective'], size=9.5, space_before=4, space_after=8,
                  color=DARK_GRAY)

    # Education
    edu_list = parse_education(data.get('education_list', []))
    if edu_list:
        section_header("EDUCATION")
        for edu in edu_list:
            p = main_para(f"{edu['degree']}", bold=True, size=10, space_before=4)
            main_para(f"{edu['institution']} | {edu['years']}" + (f" | GPA: {edu['gpa']}" if edu.get('gpa') else ''),
                      italic=True, size=9, color=MED_GRAY)

    # Work Experience
    work_list = parse_work(data.get('work_list', []))
    if work_list:
        section_header("WORK EXPERIENCE")
        for job in work_list:
            main_para(f"{job['position']}", bold=True, size=10, space_before=4)
            main_para(f"{job['company']} | {job['years']}", italic=True, size=9, color=MED_GRAY)
            if job.get('description'):
                for desc_line in job['description'].split(','):
                    main_para(f"• {desc_line.strip()}", size=9)

    # Certificates
    cert_list = parse_certificates(data.get('cert_list', []))
    if cert_list:
        section_header("CERTIFICATES")
        for cert in cert_list:
            main_para(f"• {cert['name']} – {cert['org']} ({cert['year']})", size=9.5, space_before=2)

    # Hobbies
    if data.get('hobbies'):
        section_header("INTERESTS")
        main_para(data['hobbies'], size=9.5)

    doc.save(output_path)
